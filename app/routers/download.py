import json
import logging
from io import BytesIO
from urllib.parse import quote

from fastapi import APIRouter, Depends, Query
from fastapi.responses import Response
from sqlalchemy.orm import Session

from app.database import get_db
from app.models import Resume, JobPosting
from app.services.pdf_service import PDFService

router = APIRouter(tags=["download"])
logger = logging.getLogger(__name__)

pdf_service = PDFService()


def _dispo(filename: str) -> str:
    ascii_name = filename.encode("ascii", errors="replace").decode("ascii")
    return f'attachment; filename="{ascii_name}"; filename*=UTF-8\'\'{quote(filename)}'


def _build_body_html(resume: Resume, job) -> str:
    parts = ['<div class="section">']
    name = (resume.optimized_text or "").split("\n")[0].strip() or "个人简历"
    parts.append(f"<h1>{name}</h1>")
    if job:
        parts.append(f'<p style="text-align:center;font-size:10pt;color:#666;margin-top:2px">目标: {job.position} @ {job.company_name}</p>')
    if resume.matching_score:
        parts.append(f'<div class="match-badge">匹配度: {int(resume.matching_score * 100)}%</div>')
    parts.append("</div>")

    if resume.improvements:
        try:
            improvements = json.loads(resume.improvements)
            if improvements:
                parts.append("<h2>优化说明</h2>")
                for imp in improvements:
                    parts.append(f"<p>• {imp}</p>")
        except (json.JSONDecodeError, TypeError):
            pass

    text_lines = (resume.optimized_text or "").split("\n")
    parts.append("<h2>简历详情</h2>")
    for line in text_lines[1:]:
        line = line.strip()
        if not line:
            continue
        if line.startswith("#"):
            parts.append(f"<h2>{line.lstrip('#').strip()}</h2>")
        elif line.startswith("-") or line.startswith("•"):
            parts.append(f"<p>• {line.lstrip('-• ')}</p>")
        elif line:
            parts.append(f"<p>{line}</p>")
    parts.append("</div>")
    return "\n".join(parts)


def _generate_txt(resume: Resume, job) -> bytes:
    text = resume.optimized_text or ""
    parts = []
    name = text.split("\n")[0].strip() if text else ""
    if name:
        parts.append(name)
    if job:
        parts.append(f"目标职位: {job.position} @ {job.company_name}")
    if resume.matching_score:
        parts.append(f"匹配度: {int(resume.matching_score * 100)}%")
    header = "\n".join(parts)
    if header:
        header += "\n" + "=" * 40 + "\n\n"
    return (header + text).encode("utf-8")


def _set_font(run, name="Microsoft YaHei"):
    """Set both Western and East-Asian font on a run."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)


def _generate_docx(resume: Resume, job) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    _set_font(style)
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.5

    # Document header
    text = resume.optimized_text or ""
    name = text.split("\n")[0].strip() if text else ""
    if name:
        p = doc.add_paragraph()
        run = p.add_run(name)
        _set_font(run)
        run.bold = True
        run.font.size = Pt(18)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
    if job:
        p = doc.add_paragraph()
        run = p.add_run(f"目标: {job.position} @ {job.company_name}")
        _set_font(run)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
    if resume.matching_score:
        p = doc.add_paragraph()
        run = p.add_run(f"匹配度: {int(resume.matching_score * 100)}%")
        _set_font(run)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)

    for line in (resume.optimized_text or "").split("\n"):
        line = line.strip()
        if not line:
            continue

        if line.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(line[4:])
            _set_font(run)
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
            p.paragraph_format.space_before = Pt(10)
        elif line.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(line[3:])
            _set_font(run)
            run.bold = True
            run.font.size = Pt(14)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(14)
        elif line.startswith("# "):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            _set_font(run)
            run.bold = True
            run.font.size = Pt(16)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(18)
        elif line.startswith("- ") or line.startswith("• "):
            p = doc.add_paragraph(line.lstrip("-• "), style="List Bullet")
            for run in p.runs:
                _set_font(run)
        else:
            p = doc.add_paragraph(line)
            for run in p.runs:
                _set_font(run)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


@router.get("/download/{resume_id}")
async def download_resume(
    resume_id: int,
    format: str = Query("pdf", pattern="^(pdf|txt|docx)$"),
    db: Session = Depends(get_db),
):
    resume = db.query(Resume).filter(Resume.id == resume_id).first()
    if not resume:
        return Response(status_code=404, content="简历不存在")

    job = None
    if resume.job_posting_id:
        job = db.query(JobPosting).filter(JobPosting.id == resume.job_posting_id).first()

    prefix = f"简历_{job.position}" if job else "简历"

    if format == "txt":
        content = _generate_txt(resume, job)
        return Response(
            content=content,
            media_type="text/plain; charset=utf-8",
            headers={"Content-Disposition": _dispo(f"{prefix}.txt")},
        )

    if format == "docx":
        content = _generate_docx(resume, job)
        return Response(
            content=content,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": _dispo(f"{prefix}.docx")},
        )

    title = prefix
    body_html = _build_body_html(resume, job)
    try:
        pdf_bytes = pdf_service.generate(title, body_html)
    except Exception as e:
        logger.error(f"PDF generation failed: {e}")
        return Response(
            content=f"PDF 生成失败: {str(e)}。请尝试下载 TXT 或 DOCX 格式。",
            media_type="text/plain; charset=utf-8",
        )

    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": _dispo(f"{prefix}.pdf")},
    )
